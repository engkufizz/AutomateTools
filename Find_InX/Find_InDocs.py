import os
from docx import Document

def find_string_in_docx(directory, search_string, output_file, match_whole_word=False):
    # List to store the names of files containing the search string
    files_with_string = []

    # Iterate over all files in the specified directory
    for filename in os.listdir(directory):
        # Check if the file is a .docx file
        if filename.endswith('.docx'):
            file_path = os.path.join(directory, filename)
            try:
                # Open the .docx file
                doc = Document(file_path)
                found = False

                # Iterate over each paragraph in the document
                for paragraph in doc.paragraphs:
                    paragraph_text = paragraph.text.lower()
                    search_text = search_string.lower()

                    if match_whole_word:
                        # Split the paragraph into words and check for exact matches
                        words = paragraph_text.split()
                        if search_text in words:
                            found = True
                            break
                    else:
                        if search_text in paragraph_text:
                            found = True
                            break

                # Check tables for the search string
                if not found:
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                cell_text = cell.text.lower()
                                if match_whole_word:
                                    words = cell_text.split()
                                    if search_text in words:
                                        found = True
                                        break
                                else:
                                    if search_text in cell_text:
                                        found = True
                                        break
                            if found:
                                break
                        if found:
                            break

                if found:
                    files_with_string.append(filename)

            except Exception as e:
                print(f"Error reading {filename}: {e}")

    # Write the results to the output file
    with open(output_file, 'w') as f:
        for file in files_with_string:
            f.write(file + '\n')

# Define the directory containing the .docx files
directory = r'<ENTER_DIRECTORY_PATH_HERE>'
# Define the string to search for
search_string = '<ENTER_SEARCH_STRING_HERE>'
# Define the output file
output_file = 'result.txt'

# Call the function
find_string_in_docx(directory, search_string, output_file, match_whole_word=False)

print(f"Search complete. Results saved to {output_file}.")

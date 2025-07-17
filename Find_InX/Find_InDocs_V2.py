import os
import tkinter as tk
from tkinter import filedialog
from docx import Document

def find_string_in_docx_top_directory(directory, search_string, output_file):
    """
    Searches for a string in .docx files ONLY in the top-level of a given directory.
    """
    files_with_string = []
    search_text = search_string.lower()
    print(f"Searching for '{search_string}' in .docx files within '{directory}' (no subfolders)...")

    for filename in os.listdir(directory):
        if filename.lower().endswith('.docx'):
            file_path = os.path.join(directory, filename)
            if os.path.isfile(file_path):
                if search_single_docx(file_path, search_text):
                    files_with_string.append(file_path)
    
    write_results(output_file, search_string, directory, files_with_string, recursive=False)


def find_string_in_docx_recursively(directory, search_string, output_file):
    """
    Searches for a string in .docx files in a directory and ALL its subdirectories.
    """
    files_with_string = []
    search_text = search_string.lower()
    print(f"Searching for '{search_string}' in .docx files within '{directory}' and all subfolders...")

    for dirpath, _, filenames in os.walk(directory):
        for filename in filenames:
            if filename.lower().endswith('.docx'):
                file_path = os.path.join(dirpath, filename)
                if search_single_docx(file_path, search_text):
                    files_with_string.append(file_path)

    write_results(output_file, search_string, directory, files_with_string, recursive=True)


def search_single_docx(file_path, search_text):
    """
    Checks a single .docx file for the presence of search_text.
    Returns True if found, False otherwise.
    """
    try:
        doc = Document(file_path)
        # Search in paragraphs
        for para in doc.paragraphs:
            if search_text in para.text.lower():
                return True
        
        # Search in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if search_text in cell.text.lower():
                        return True
    except Exception as e:
        print(f"Could not read file {os.path.basename(file_path)} due to error: {e}")
    
    return False


def write_results(output_file, search_string, directory, file_list, recursive):
    """Writes the list of found files to the output file."""
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            if file_list:
                f.write(f"Files containing the string '{search_string}':\n\n")
                for file_path in file_list:
                    f.write(file_path + '\n')
            else:
                search_area = f"'{directory}' and its subdirectories" if recursive else f"'{directory}'"
                f.write(f"No .docx files found in {search_area} containing the string '{search_string}'.\n")
        print(f"\nSearch complete. Results saved to {output_file}.")
    except Exception as e:
        print(f"Error writing to output file {output_file}: {e}")


# --- Main Script Execution ---
if __name__ == "__main__":
    # 1. Set up and hide the root Tkinter window
    root = tk.Tk()
    root.withdraw()

    # 2. Prompt user to select a directory
    print("A folder selection window will now open.")
    directory = filedialog.askdirectory(title="Select the folder to search for .docx files")

    # 3. Proceed only if a directory was selected
    if directory:
        print(f"Directory selected: {directory}\n")

        # 4. Ask user whether to include subfolders
        include_subfolders_choice = input("Include subfolders in the search? (yes/no): ").lower().strip()
        
        # 5. Get the search string from the user
        search_string = input("Please enter the string to search for: ")

        if search_string:
            output_file = 'result.txt'
            
            # 6. Call the appropriate function based on user's choice
            if include_subfolders_choice.startswith('y'):
                find_string_in_docx_recursively(directory, search_string, output_file)
            else:
                find_string_in_docx_top_directory(directory, search_string, output_file)
        else:
            print("No search string entered. Exiting programme.")
    else:
        print("No directory was selected. Exiting programme.")

    # Keep the console window open until the user presses Enter
    input("\nPress Enter to exit...")
